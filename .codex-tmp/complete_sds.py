from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(sys.argv[1])
TARGET = Path(sys.argv[2])
DOWNLOADS = Path(r"C:\Users\ontri\Downloads")
ROOT = Path(r"D:\SWD\UNI-MATE")
GENERATED = ROOT / ".codex-tmp" / "sds-doc" / "diagrams"
OPTIMIZED = ROOT / ".codex-tmp" / "optimized-diagrams-v1"

# The supplied draw.io exports are trusted, very high-resolution PNGs.
Image.MAX_IMAGE_PIXELS = None

ARCHITECTURE = OPTIMIZED / "architecture.png"
PACKAGE = OPTIMIZED / "package.png"
USE_CASE = DOWNLOADS / "UNI-MATE_Use_Case_Diagram_Corrected.png"
CONCEPTUAL = OPTIMIZED / "conceptual.png"
CLASS_DIAGRAM = OPTIMIZED / "class.png"
ACTIVITY = OPTIMIZED / "activity.png"
SEQ_REGISTER = OPTIMIZED / "seq01_registration.png"
SEQ_LOGIN = OPTIMIZED / "seq02_login.png"
SEQ_FORGOT = OPTIMIZED / "seq03_forgot.png"
SEQ_CHAT = OPTIMIZED / "seq06_chat.png"

RED = "C00000"
DARK_RED = "8B2E1A"
PEACH = "FCE4D6"
LIGHT_RED = "F4CCCC"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9E1F2"
GRAY = "666666"


def set_run_font(run, name="Calibri", size=10.5, bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def find_paragraph(doc, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def first_drawing_after(doc, paragraph):
    body_children = list(doc.element.body)
    start = body_children.index(paragraph._p)
    for child in body_children[start + 1 :]:
        if child.tag == qn("w:p") and child.xpath(".//w:drawing"):
            from docx.text.paragraph import Paragraph

            return Paragraph(child, doc)
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            candidate = Paragraph(child, doc)
            if candidate.text.strip() and candidate.style.name.startswith("Heading"):
                break
    raise ValueError(f"No drawing found after {paragraph.text}")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_paragraph_before(doc, reference, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    reference._p.addprevious(paragraph._p)
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_heading_before(doc, reference, text: str, level: int, page_break=False):
    paragraph = insert_paragraph_before(doc, reference, text, style=f"Heading {level}")
    paragraph.paragraph_format.page_break_before = page_break
    return paragraph


def insert_label_before(doc, reference, text: str, page_break=False):
    paragraph = insert_paragraph_before(doc, reference)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=True, color=DARK_RED)
    return paragraph


def insert_body_before(doc, reference, text: str, *, italic=False, color=None):
    paragraph = insert_paragraph_before(doc, reference)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, italic=italic, color=color)
    return paragraph


def best_size(path: Path, max_width=6.35, max_height=6.65):
    with Image.open(path) as image:
        width_px, height_px = image.size
    width = max_width
    height = width * height_px / width_px
    if height > max_height:
        height = max_height
        width = height * width_px / height_px
    return width, height


def add_picture_to_paragraph(paragraph, path: Path, *, max_width=6.35, max_height=6.65, crop=None):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True

    if crop is None:
        width, height = best_size(path, max_width, max_height)
    else:
        left, top, right, bottom = crop
        with Image.open(path) as image:
            width_px, height_px = image.size
        crop_width = width_px * (right - left)
        crop_height = height_px * (bottom - top)
        width = max_width
        height = width * crop_height / crop_width
        if height > max_height:
            height = max_height
            width = height * crop_width / crop_height

    inline = paragraph.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))

    if crop is not None:
        left, top, right, bottom = crop
        blip_fill = inline._inline.find(".//" + qn("pic:blipFill"))
        src_rect = blip_fill.find(qn("a:srcRect"))
        if src_rect is None:
            src_rect = OxmlElement("a:srcRect")
            blip = blip_fill.find(qn("a:blip"))
            blip.addnext(src_rect)
        src_rect.set("l", str(round(left * 100000)))
        src_rect.set("t", str(round(top * 100000)))
        src_rect.set("r", str(round((1 - right) * 100000)))
        src_rect.set("b", str(round((1 - bottom) * 100000)))
    return paragraph


def insert_picture_before(doc, reference, path: Path, *, max_width=6.35, max_height=6.65, crop=None):
    paragraph = insert_paragraph_before(doc, reference)
    return add_picture_to_paragraph(
        paragraph,
        path,
        max_width=max_width,
        max_height=max_height,
        crop=crop,
    )


def insert_caption_before(doc, reference, text: str):
    paragraph = insert_paragraph_before(doc, reference)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9, italic=True, color=GRAY)
    return paragraph


def set_cell_margins(cell, top=130, start=150, bottom=130, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_inches=6.4):
    width_dxa = int(width_inches * 1440)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(width_dxa))
    grid.append(grid_col)
    cell = table.cell(0, 0)
    tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def insert_callout_before(doc, reference, title: str, text: str, *, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    reference._p.addprevious(table._tbl)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_run_font(run, size=10, bold=True, color=DARK_RED)
    body = cell.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_after = Pt(0)
    run = body.add_run(text)
    set_run_font(run, size=9.5)
    spacer = insert_paragraph_before(doc, reference)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def insert_placeholder_before(doc, reference, title: str, owner: str, scope: str, *, page_break=False):
    insert_label_before(doc, reference, title, page_break=page_break)
    table = doc.add_table(rows=1, cols=1)
    reference._p.addprevious(table._tbl)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=260, start=180, bottom=260, end=180)
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("PLACEHOLDER - UML SEQUENCE DIAGRAM NOT YET SUPPLIED")
    set_run_font(run, size=11, bold=True, color=DARK_RED)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Owner: {owner}")
    set_run_font(run, size=10, bold=True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(scope)
    set_run_font(run, size=9.5)
    insert_caption_before(doc, reference, f"{title}. Reserved position for the final team-authored UML sequence diagram")


def remove_between(start_paragraph, end_paragraph):
    current = start_paragraph._p.getnext()
    while current is not None and current is not end_paragraph._p:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def replace_caption(doc, old_text: str, new_text: str):
    paragraph = find_paragraph(doc, old_text)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    return paragraph


def add_sequence_image(doc, reference, label: str, path: Path, caption: str, *, page_break=False, panels=None):
    insert_label_before(doc, reference, label, page_break=page_break)
    if panels:
        for idx, crop in enumerate(panels, start=1):
            insert_picture_before(doc, reference, path, max_width=6.2, max_height=6.15, crop=crop)
            insert_caption_before(doc, reference, f"{caption} - part {idx} of {len(panels)}")
    else:
        insert_picture_before(doc, reference, path, max_width=6.2, max_height=6.55)
        insert_caption_before(doc, reference, caption)


def update_change_record(doc):
    table = doc.tables[0]
    source_row = deepcopy(table.rows[2]._tr)
    target_row = table.rows[-1]._tr
    target_row.addprevious(source_row)
    target_row.getparent().remove(target_row)
    row = table.rows[-1]
    values = [
        "21/07/2026",
        "M",
        "UNI-MATE Team",
        "Integrated the supplied architecture, package, use-case, class, conceptual, activity, and sequence diagrams; corrected diagram sizing/pagination; reserved explicit positions for missing SD-04 and SD-08; and documented the current 2FA implementation gap.",
    ]
    for cell, value in zip(row.cells, values):
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(value)


def main() -> int:
    for path in (
        ARCHITECTURE,
        PACKAGE,
        USE_CASE,
        CONCEPTUAL,
        CLASS_DIAGRAM,
        ACTIVITY,
        SEQ_REGISTER,
        SEQ_LOGIN,
        SEQ_FORGOT,
        SEQ_CHAT,
    ):
        if not path.exists():
            raise FileNotFoundError(path)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, TARGET)
    doc = Document(TARGET)
    update_change_record(doc)

    # Replace the top-level architecture and package images with the team-supplied diagrams.
    architecture_heading = find_paragraph(doc, "a. Architectural View")
    add_picture_to_paragraph(first_drawing_after(doc, architecture_heading), ARCHITECTURE, max_width=6.35, max_height=4.1)
    package_heading = find_paragraph(doc, "b. Package Dependency Diagram")
    add_picture_to_paragraph(first_drawing_after(doc, package_heading), PACKAGE, max_width=5.25, max_height=6.7)

    # Add the requested conceptual and design class diagrams while preserving the implementation schemas.
    core_heading = find_paragraph(doc, "a. Core Matching and Chat Schema")
    insert_heading_before(doc, core_heading, "a. Conceptual Domain Class Diagram", 3)
    insert_picture_before(doc, core_heading, CONCEPTUAL, max_width=6.35, max_height=4.0)
    insert_caption_before(doc, core_heading, "UML Figure U1. UNI-MATE conceptual domain model and multiplicities")
    insert_body_before(
        doc,
        core_heading,
        "This is a conceptual, class-style domain view. The collection descriptions and schema constraints below remain authoritative for persistence details, indexes, embedded values, and implementation-specific cardinality rules.",
        italic=True,
        color=GRAY,
    )
    insert_heading_before(doc, core_heading, "b. Design Class Diagram", 3)
    insert_picture_before(doc, core_heading, CLASS_DIAGRAM, max_width=6.35, max_height=3.3)
    insert_caption_before(doc, core_heading, "UML Figure U2. Core account, matching, chat, notification, and safety classes")
    core_heading.text = "c. Core Matching and Chat Schema"
    find_paragraph(doc, "b. Extended Social, Partner, and Moderation Schema").text = "d. Extended Social, Partner, and Moderation Schema"
    find_paragraph(doc, "c. Collection Descriptions").text = "e. Collection Descriptions"
    find_paragraph(doc, "d. Important Indexes and Constraints").text = "f. Important Indexes and Constraints"

    # Add the use-case view and five activity diagrams before Code Designs.
    code_designs = find_paragraph(doc, "II. Code Designs")
    insert_heading_before(doc, code_designs, "4. UML Functional and Behavioral Views", 2, page_break=False)
    insert_body_before(
        doc,
        code_designs,
        "This section consolidates the team-supplied UML functional and activity views. Diagrams are retained as authored; the alignment notes distinguish current repository behavior from planned scope so the SDS remains an as-is design reference.",
    )
    insert_heading_before(doc, code_designs, "a. Use Case Diagram", 3)
    insert_picture_before(doc, code_designs, USE_CASE, max_width=6.35, max_height=4.45)
    insert_caption_before(doc, code_designs, "UML Figure U3. UNI-MATE use-case model")
    insert_callout_before(
        doc,
        code_designs,
        "Repository alignment",
        "The supplied use-case view includes registration OTP, exported growth reports, and coffee-shop report export as broader/planned scope. In the current repository, registration sets emailVerified directly, EmailOtp is used only for forgot-password, and admin analytics exposes dashboard counts rather than CSV/Excel export.",
        fill=LIGHT_RED,
    )
    insert_heading_before(doc, code_designs, "b. Activity Diagrams", 3, page_break=True)
    activity_blocks = [
        ("Activity AD-01. User Registration and Profile Setup", "UML Figure U4. User registration and profile setup activity", (0.00, 0.00, 0.43, 0.37)),
        ("Activity AD-02. Discovery, Matching and Chatting", "UML Figure U5. Discovery, matching, and chatting activity", (0.49, 0.00, 1.00, 0.37)),
        ("Activity AD-03. Report User", "UML Figure U6. User report and moderation activity", (0.00, 0.385, 0.46, 0.70)),
        ("Activity AD-04. Group Study", "UML Figure U7. Group study creation and chat activity", (0.48, 0.395, 1.00, 0.705)),
        ("Activity AD-05. Cafe Registration and Admin Approval", "UML Figure U8. Cafe registration and admin approval activity", (0.015, 0.72, 0.54, 1.00)),
    ]
    for idx, (label, caption, crop) in enumerate(activity_blocks):
        insert_label_before(doc, code_designs, label, page_break=idx > 0)
        activity_height = 5.35 if idx == 3 else 6.1
        insert_picture_before(doc, code_designs, ACTIVITY, max_width=6.1, max_height=activity_height, crop=crop)
        insert_caption_before(doc, code_designs, caption)
    insert_callout_before(
        doc,
        code_designs,
        "As-is behavior note",
        "The activity views use UML swimlanes, actions, decisions, and final nodes. The current mutual-like implementation opens an active ChatRoom immediately, even though the product README describes cafe-gated chat; Appendix C records this conflict.",
        fill=LIGHT_RED,
    )
    code_designs.paragraph_format.page_break_before = True

    # Rebuild the authentication sequence section so all supplied diagrams actually render.
    auth_sequence = find_paragraph(doc, "c. Sequence Diagram(s)")
    auth_db = find_paragraph(doc, "d. Database Queries / Persistence Operations")
    remove_between(auth_sequence, auth_db)
    auth_sequence.paragraph_format.page_break_before = True
    add_sequence_image(
        doc,
        auth_db,
        "SD-01. Đăng ký tài khoản (Luân)",
        SEQ_REGISTER,
        "Sequence Diagram SD-01. Account registration",
    )
    add_sequence_image(
        doc,
        auth_db,
        "SD-02. Đăng nhập và 2FA (Luân)",
        SEQ_LOGIN,
        "Sequence Diagram SD-02. Current password-login flow",
        page_break=True,
        panels=[(0.0, 0.0, 1.0, 0.50), (0.0, 0.50, 1.0, 1.0)],
    )
    insert_callout_before(
        doc,
        auth_db,
        "2FA implementation gap",
        "The supplied login diagram correctly reflects the current code path: loginWithPassword issues access and refresh tokens immediately after password validation. No active login OTP/2FA branch exists for user, partner, or admin accounts. README and /admin/settings still claim admin 2FA is required; this mismatch is retained in Appendix C.",
        fill=LIGHT_RED,
    )
    add_sequence_image(
        doc,
        auth_db,
        "SD-03. Quên mật khẩu (Luân)",
        SEQ_FORGOT,
        "Sequence Diagram SD-03. Forgot-password OTP and reset-token flow",
        page_break=False,
        panels=[(0.0, 0.0, 1.0, 0.50), (0.0, 0.50, 1.0, 1.0)],
    )
    insert_placeholder_before(
        doc,
        auth_db,
        "SD-04. Onboarding và vị trí",
        "Khang",
        "Expected scope: OnboardingPage -> userApi/Axios -> users.router/Zod -> UserController -> User model, including age validation, GPS/manual location selection, geocoding, and onboardingCompleted persistence.",
        page_break=True,
    )

    # Number the available implementation-aligned sequence diagrams.
    replace_caption(
        doc,
        "Figure 8. Discovery like/pass decision and actual mutual-match behavior",
        "Sequence Diagram SD-05. Discovery, like/pass, and match creation (Khang)",
    )

    # Replace private chat with the supplied detailed realtime diagram.
    chat_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "c. Sequence Diagram(s)" and paragraph._p.getprevious() is not None:
            # Locate the occurrence belonging to the private chat section via following caption.
            current = paragraph._p.getnext()
            while current is not None:
                if current.tag == qn("w:p"):
                    text = "".join(current.itertext())
                    if "Figure 12. Realtime private message" in text:
                        chat_heading = paragraph
                        break
                    if text.strip().startswith("d. Database Queries"):
                        break
                current = current.getnext()
        if chat_heading is not None:
            break
    if chat_heading is None:
        raise ValueError("Private-chat sequence section not found")
    chat_db = None
    current = chat_heading._p.getnext()
    from docx.text.paragraph import Paragraph

    while current is not None:
        if current.tag == qn("w:p"):
            candidate = Paragraph(current, doc)
            if candidate.text.strip() == "d. Database Queries / Persistence Operations":
                chat_db = candidate
                break
        current = current.getnext()
    remove_between(chat_heading, chat_db)
    add_sequence_image(
        doc,
        chat_db,
        "SD-06. Chat cá nhân realtime (Duy)",
        SEQ_CHAT,
        "Sequence Diagram SD-06. Realtime private chat, block-state checks, notification, and delivery flow",
    )

    replace_caption(
        doc,
        "Figure 14. Group creation, member addition, and realtime group messaging",
        "Sequence Diagram SD-07. Group management and group chat (Toàn)",
    )

    # Safety: reserve the missing block/unblock sequence and retain the correct report/admin sequence.
    safety_caption = replace_caption(
        doc,
        "Figure 18. Report submission and admin moderation decision",
        "Sequence Diagram SD-09. Admin report review and audited moderation action (Phúc)",
    )
    safety_image = safety_caption._p.getprevious()
    from docx.text.paragraph import Paragraph

    safety_reference = Paragraph(safety_image, doc)
    insert_placeholder_before(
        doc,
        safety_reference,
        "SD-08. Báo cáo, block và unblock",
        "Toàn",
        "Expected scope: User/SafetyPage -> SafetyController -> Report/User/Match/ChatRoom -> Socket/Notification. The final diagram must show contextual report creation, bilateral block effects, and unblock restoration only when the peer is not still blocking.",
        page_break=False,
    )

    # Split the combined partner/voucher sequence into SD-10 and SD-11 panels.
    places_caption = find_paragraph(doc, "Figure 16. Partner place approval, voucher creation, and user save flow")
    places_image = Paragraph(places_caption._p.getprevious(), doc)
    add_picture_to_paragraph(places_image, GENERATED / "15_seq_partner_voucher.png", max_width=6.2, max_height=5.4, crop=(0.0, 0.0, 1.0, 0.48))
    places_caption.text = "Sequence Diagram SD-10. Partner cafe registration and admin approval (Hoàng)"
    places_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if places_caption.runs:
        set_run_font(places_caption.runs[0], size=9, italic=True, color=GRAY)
    insert_label_before(doc, places_image, "SD-10. Partner đăng ký quán và admin duyệt (Hoàng)")
    places_db = None
    current = places_caption._p.getnext()
    while current is not None:
        if current.tag == qn("w:p"):
            candidate = Paragraph(current, doc)
            if candidate.text.strip() == "d. Database Queries / Persistence Operations":
                places_db = candidate
                break
        current = current.getnext()
    insert_label_before(doc, places_db, "SD-11. Tạo và lưu voucher (Hoàng)", page_break=True)
    insert_picture_before(doc, places_db, GENERATED / "15_seq_partner_voucher.png", max_width=6.2, max_height=5.5, crop=(0.0, 0.43, 1.0, 1.0))
    insert_caption_before(doc, places_db, "Sequence Diagram SD-11. Partner creates a voucher and a user saves it")

    # The Admin section references SD-09 instead of duplicating the same sequence image.
    admin_caption = find_paragraph(doc, "Figure 20. Representative admin review and audited moderation flow")
    admin_image = admin_caption._p.getprevious()
    admin_image.getparent().remove(admin_image)
    admin_caption.text = "See Sequence Diagram SD-09 in Section II.7 for the admin report review and audited moderation flow."
    admin_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if admin_caption.runs:
        set_run_font(admin_caption.runs[0], size=10.5, italic=True, color=GRAY)

    # Add a concise diagram register after the Code Designs introduction.
    code_intro = find_paragraph(
        doc,
        "This section groups classes and runtime interactions by user-visible function. Boundary objects are React pages, control objects are controllers/services/socket handlers, and entity objects are Mongoose models.",
    )
    next_element = code_intro._p.getnext()
    next_paragraph = Paragraph(next_element, doc)
    insert_callout_before(
        doc,
        next_paragraph,
        "Sequence diagram register",
        "SD-01 Registration (Luân); SD-02 Login/2FA scope (Luân); SD-03 Forgot password (Luân); SD-04 Onboarding/location (Khang, placeholder); SD-05 Discovery/match (Khang); SD-06 Private realtime chat (Duy); SD-07 Group/chat (Toàn); SD-08 Report/block/unblock (Toàn, placeholder); SD-09 Admin report handling (Phúc); SD-10 Partner cafe approval (Hoàng); SD-11 Voucher create/save (Hoàng).",
        fill=MID_GRAY,
    )

    # Update the static TOC with the added UML section; page numbers are corrected after render QA.
    toc_anchor = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name.lower() == "toc 1" and paragraph.text.strip().startswith("II. Code Designs")
    )
    toc_source = toc_anchor._p.getprevious()
    toc_clone = deepcopy(toc_source)
    toc_anchor._p.addprevious(toc_clone)
    toc_line = Paragraph(toc_clone, doc)
    clear_paragraph(toc_line)
    toc_line.add_run("4. UML Functional and Behavioral Views")
    toc_line.paragraph_format.space_after = Pt(0)

    toc_pages = {
        "I. Overview": 4,
        "1. System Architecture and Code Packages": 4,
        "2. Design Conventions": 5,
        "3. Database Design": 6,
        "4. UML Functional and Behavioral Views": 11,
        "II. Code Designs": 17,
        "1. Authentication, Account, and Onboarding": 18,
        "2. Discovery and Mutual Matching": 25,
        "3. Cafe Proposal and Match Lifecycle": 27,
        "4. Private Chat and Notifications": 29,
        "5. Groups and Group Chat": 31,
        "6. Places, Partner Registration, and Vouchers": 33,
        "7. Safety, Blocking, Reporting, and Moderation": 35,
        "8. Administration and Operational Oversight": 37,
        "III. Cross-Cutting Design": 39,
        "1. Security and Authorization": 39,
        "2. Realtime Design": 39,
        "3. State Models": 40,
        "4. External Services and Failure Handling": 40,
        "Appendix A. REST API Inventory": 41,
        "Appendix B. Configuration and Deployment": 43,
        "1. Runtime Stack": 43,
        "2. Required Configuration Categories": 43,
        "3. Build and Run": 43,
        "Appendix C. Known Implementation Discrepancies and Risks": 44,
    }
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.lower().startswith("toc"):
            continue
        label = paragraph.text.rsplit("\t", 1)[0].strip()
        if label not in toc_pages:
            continue
        clear_paragraph(paragraph)
        paragraph.add_run(f"{label}\t{toc_pages[label]}")

    doc.save(TARGET)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
