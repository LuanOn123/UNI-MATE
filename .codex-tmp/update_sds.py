from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(sys.argv[1])
TARGET = Path(sys.argv[2])


def replace_paragraph_text(doc, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text == old:
            if paragraph.runs:
                paragraph.runs[0].text = new
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(new)
            return
    raise ValueError(f"Paragraph not found: {old}")


def set_cell_text(cell, value: str) -> None:
    paragraphs = cell.paragraphs
    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = value
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(value)
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def clone_last_row(table):
    table._tbl.append(deepcopy(table.rows[-1]._tr))
    return table.rows[-1]


def remove_row(table, index: int) -> None:
    row = table.rows[index]
    row._tr.getparent().remove(row._tr)


def replace_row_with_clone(table, source_index: int, target_index: int) -> None:
    source = deepcopy(table.rows[source_index]._tr)
    target = table.rows[target_index]._tr
    parent = target.getparent()
    position = list(parent).index(target)
    parent.remove(target)
    parent.insert(position, source)


def set_row(table, index: int, values: list[str]) -> None:
    row = table.rows[index]
    if len(row.cells) != len(values):
        raise ValueError(f"Row {index} has {len(row.cells)} cells, expected {len(values)}")
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def restart_numbering(doc, paragraph_texts: list[str]) -> None:
    paragraphs = []
    for text in paragraph_texts:
        paragraph = next((p for p in doc.paragraphs if p.text == text), None)
        if paragraph is None:
            raise ValueError(f"Numbered paragraph not found: {text}")
        paragraphs.append(paragraph)

    source_num_id = int(paragraphs[0]._p.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    source_num = next(num for num in numbering.xpath('./w:num') if int(num.get(qn('w:numId'))) == source_num_id)
    abstract_num_id = source_num.find(qn('w:abstractNumId')).get(qn('w:val'))
    new_num_id = max(int(num.get(qn('w:numId'))) for num in numbering.xpath('./w:num')) + 1

    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), str(new_num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), abstract_num_id)
    new_num.append(abstract_ref)
    level_override = OxmlElement('w:lvlOverride')
    level_override.set(qn('w:ilvl'), '0')
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)

    for paragraph in paragraphs:
        paragraph._p.pPr.numPr.numId.val = new_num_id


def main() -> int:
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, TARGET)
    doc = Document(TARGET)

    # Record of changes and implementation baseline.
    replace_row_with_clone(doc.tables[0], 1, 2)
    set_row(doc.tables[0], 2, [
        "21/07/2026",
        "M",
        "UNI-MATE Team",
        "Reconciled the SDS with the current repository, especially direct password login, password-reset OTP, API inventory, security controls, and implementation risks. Diagram images were intentionally left unchanged.",
    ])
    set_cell_text(
        doc.tables[1].cell(0, 0),
        "Implementation baseline: This SDS documents the code present in the repository on 21 July 2026. It describes actual behavior, including direct password login for every role, password-reset OTP, the current mutual-like chat behavior, and the remaining cafe proposal workflow.",
    )

    # Authentication and OTP scope.
    set_cell_text(
        doc.tables[3].cell(2, 2),
        "Hashed email OTP used by the forgot-password flow, with expiry, attempt count, resend count, consumed flag, and TTL cleanup.",
    )
    set_cell_text(
        doc.tables[2].cell(4, 2),
        "Zustand authentication/session state and persisted access/refresh token context.",
    )
    set_cell_text(
        doc.tables[3].cell(16, 2),
        "Standalone cafe/profile tag catalogue model; no active admin route or controller manages it, and places persist tag names as strings.",
    )
    replace_paragraph_text(
        doc,
        "This subsystem creates accounts, validates credentials, optionally performs email OTP two-factor authentication, refreshes sessions, resets passwords, and collects the profile/preferences/location required for discovery.",
        "This subsystem creates accounts, authenticates every role directly with email and password, refreshes sessions, resets forgotten passwords through an email OTP plus a short-lived reset token, and collects the profile, preferences, and location required for discovery.",
    )
    replace_paragraph_text(
        doc,
        "Figure 6. Password login with optional email OTP two-factor authentication",
        "Figure 6. Direct password login and the email-OTP password-reset flow",
    )

    auth_rows = [
        ["Component", "Method", "Inputs / outputs and internal processing"],
        ["AuthController", "registerController", "email, password -> user/tokens; delegates registration, sets the refresh cookie, and returns the new session."],
        ["AuthController", "loginController", "email, password -> user/tokens for every role; sets the refresh cookie without an OTP or 2FA branch."],
        ["AuthController", "forgotPasswordSendOtpController", "email -> generic OTP-sent response after the password-reset OTP service succeeds."],
        ["AuthController", "forgotPasswordVerifyOtpController", "email, six-digit OTP -> short-lived resetToken after OTP validation and consumption."],
        ["AuthController", "resetPasswordController", "resetToken, newPassword -> password replacement; rejects reuse of the current password."],
        ["AuthService", "registerWithPassword", "Normalizes email, rejects duplicates, hashes the password, creates an emailVerified user, and issues tokens."],
        ["AuthService", "loginWithPassword", "Restores expired suspension, checks account state and password, updates login timestamps, and immediately issues tokens."],
        ["AuthService", "sendPasswordResetOtp / verifyPasswordResetOtp", "Requires an existing password account, stores/verifies the latest EmailOtp, consumes success, and signs a 10-minute password-reset token."],
        ["AuthService", "resetPasswordWithToken", "Validates token purpose and account state, rejects the old password, hashes the replacement, and clears the refresh-token hash."],
        ["UserController", "completeOnboarding / geocodeLocation / updateProfile", "Completes the age-gated profile, provides Vietnam-filtered manual geocoding, and persists whitelisted profile/preference updates."],
    ]
    auth_table = doc.tables[5]
    while len(auth_table.rows) < len(auth_rows):
        clone_last_row(auth_table)
    while len(auth_table.rows) > len(auth_rows):
        remove_row(auth_table, len(auth_table.rows) - 1)
    for idx, values in enumerate(auth_rows):
        set_row(auth_table, idx, values)

    set_cell_text(
        doc.tables[6].cell(0, 0),
        'const user = await User.findOne({ email });\n'
        'const tokens = await issueTokens(user); // direct password login for all roles\n'
        'await EmailOtp.create({ email, otpHash, expiresAt }); // forgot-password only\n'
        'const otp = await EmailOtp.findOne({ email, consumed: false, expiresAt: { $gt: new Date() } }).sort({ createdAt: -1 });\n'
        'user.passwordHash = await bcrypt.hash(newPassword, 10); user.refreshTokenHash = undefined; await user.save();\n'
        'await User.findByIdAndUpdate(userId, { onboardingCompleted: true, onboarding, location });',
    )

    # Security table: current login no longer has 2FA.
    set_row(doc.tables[22], 1, [
        "Authentication",
        "Short-lived JWT access token and refresh token hash/cookie. Registration and login issue tokens directly after password validation for user, partner, and admin roles.",
    ])
    set_row(doc.tables[22], 3, [
        "Password/OTP",
        "bcrypt password and OTP hashes. EmailOtp is expiring, consumable, resend/attempt-limited, and currently used only for forgot-password; login has no OTP/2FA enforcement.",
    ])

    # REST API inventory: general login OTP endpoints were removed.
    api_table = doc.tables[26]
    remove_row(api_table, 4)
    remove_row(api_table, 3)
    set_cell_text(api_table.cell(1, 3), "Create account, mark email verified, and issue tokens.")
    set_cell_text(api_table.cell(2, 3), "Validate password and issue tokens directly for every role.")

    # Configuration details.
    set_row(doc.tables[27], 2, [
        "Authentication",
        "JWT access/refresh secrets and token lifetimes; cookie secure flag. The access secret also signs 10-minute password-reset tokens.",
    ])
    set_row(doc.tables[27], 4, [
        "Email",
        "SMTP host, port, credentials, sender, and development OTP logging for the forgot-password flow.",
    ])

    replace_paragraph_text(
        doc,
        "All routes below are prefixed with /api. Except for register/login/OTP/refresh, routes require an authenticated access token. Admin and partner groups apply additional role checks.",
        "All routes below are prefixed with /api. Except for register, login, forgot-password OTP/reset, and refresh, routes require an authenticated access token. Admin and partner groups apply additional role checks.",
    )
    set_row(doc.tables[25], 4, [
        "SMTP/email",
        "Forgot-password OTP delivery.",
        "Development may print OTP when configured; service errors propagate through the centralized handler.",
    ])

    restart_numbering(doc, [
        "Frontend: React, TypeScript, Vite, React Router, Zustand, Axios, Tailwind CSS, Socket.IO Client.",
        "Backend: Node.js, Express, TypeScript, Mongoose, JWT, bcryptjs, Zod, Socket.IO, Multer, Nodemailer.",
        "Database: MongoDB with geospatial and TTL indexes.",
        "Development entry points: frontend http://localhost:5173 and backend http://localhost:5000.",
    ])
    restart_numbering(doc, [
        "Provide backend and frontend environment files.",
        "Install backend dependencies and start MongoDB or configure MongoDB Atlas.",
        "Run the backend development process.",
        "Install frontend dependencies and run the Vite client.",
        "Verify /health, authentication, protected REST routes, and Socket.IO connectivity.",
    ])

    # Current implementation discrepancies and risks.
    risk_table = doc.tables[28]
    set_row(risk_table, 2, [
        "High",
        "Cafe state is overwritten by read endpoints",
        "selectPlace sets cafe_proposed, but listMatches/getMatch call openChatForMatch and can immediately restore chat_opened. If that happens before confirmation, the active proposal may remain unaccepted and the room may not receive the selected cafe.",
    ])
    set_row(risk_table, 3, [
        "High",
        "Admin 2FA contract mismatch",
        "Current login issues tokens directly for admin accounts, the general /auth/send-otp and /auth/verify-otp endpoints and User.twoFactorEnabled field are removed, while README and /admin/settings still claim admin 2FA is required.",
    ])
    set_row(risk_table, 4, [
        "Medium",
        "Registration verification",
        "Registration sets emailVerified=true without an OTP. EmailOtp is currently used only for forgot-password and does not verify registration or login.",
    ])
    set_row(risk_table, 5, [
        "Medium",
        "Match participant cardinality",
        "Match.users is an unrestricted array; the schema does not enforce exactly two users.",
    ])
    set_row(risk_table, 6, [
        "Medium",
        "Tag management and consistency",
        "Tag has its own catalogue model but no active admin route/controller manages it; PlaceCache stores tag names as strings and does not reference Tag ids.",
    ])
    set_row(risk_table, 7, [
        "Medium",
        "Unblock restoration",
        "Unblock restores all blocked pair matches to chat_opened, even if their prior state was not an open chat.",
    ])
    set_row(risk_table, 8, [
        "Low",
        "Dual OTP models",
        "Both Otp and EmailOtp exist; the current auth service uses EmailOtp only for forgot-password, while Otp is unused by active routes.",
    ])
    if len(risk_table.rows) < 10:
        clone_last_row(risk_table)
    set_row(risk_table, 9, [
        "Low",
        "Local file storage",
        "Uploads are stored on the application host and require persistent/shared storage for scaled deployment.",
    ])

    recommendation = doc.tables[29].cell(0, 0).paragraphs[0]
    recommendation.runs[0].text = "Recommended decisions:"
    recommendation.runs[1].text = " Either enforce cafe confirmation before creating ChatRoom, or formally remove the cafe-gated requirement and stop moving chat_opened matches back to cafe_proposed. Separately, either restore enforced admin 2FA or update README and /admin/settings so the published policy matches the login implementation."
    for run in recommendation.runs[2:]:
        run.text = ""

    doc.save(TARGET)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
