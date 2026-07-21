from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def iter_blocks(parent):
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def main() -> int:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)
    records = []
    paragraph_index = 0
    table_index = 0
    for block_index, block in enumerate(iter_blocks(doc)):
        if isinstance(block, Paragraph):
            records.append({
                "block": block_index,
                "kind": "paragraph",
                "paragraph": paragraph_index,
                "style": block.style.name if block.style else None,
                "text": block.text,
                "has_drawing": bool(block._p.xpath('.//w:drawing | .//w:pict')),
            })
            paragraph_index += 1
        else:
            rows = []
            for row in block.rows:
                rows.append([cell.text for cell in row.cells])
            records.append({
                "block": block_index,
                "kind": "table",
                "table": table_index,
                "style": block.style.name if block.style else None,
                "rows": rows,
            })
            table_index += 1

    payload = {
        "source": str(source),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "records": records,
    }
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
