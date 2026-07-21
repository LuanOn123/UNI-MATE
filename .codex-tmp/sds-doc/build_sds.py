from pathlib import Path
from zipfile import ZipFile
import shutil

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"D:\SWD\UNI-MATE")
TMP = ROOT / ".codex-tmp" / "sds-doc"
REFERENCE = Path(r"C:\Users\ontri\Downloads\Template3_SDS Document.docx")
OUTPUT = ROOT / "docs" / "UNI-MATE_Software_Design_Specification.docx"
DIAGRAMS = TMP / "diagrams"
LOGO = TMP / "fpt-logo.png"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(REFERENCE, OUTPUT)
with ZipFile(REFERENCE) as zf:
    LOGO.write_bytes(zf.read("word/media/image1.png"))

doc = Document(OUTPUT)

# Remove reference placeholder body while retaining its styles, theme, numbering,
# section settings, headers, footers, and media package.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.81)
section.different_first_page_header_footer = True

RED = "C00000"
DARK_RED = "8B2E1A"
PEACH = "FCE4D6"
LIGHT_BLUE = "DDEBF7"
LIGHT_GRAY = "F2F2F2"
DARK_GREEN = "0B3D0B"
BLUE = "2E75B6"
GRAY = "666666"


def set_run_font(run, name="Calibri", size=10.5, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_fixed_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_keep(paragraph, next_=False, lines=False):
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = lines


def add_field(paragraph, code, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def create_numbering(num_fmt, level_text):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "240")
    p_pr.extend([tabs, ind])
    lvl.extend([start, fmt, text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    # Word may otherwise merge adjacent list definitions during a save/update.
    # An explicit start override guarantees that every generated list starts at 1.
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_styles():
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 20, RED, 12, 8),
        ("Heading 2", 15, "111111", 10, 5),
        ("Heading 3", 12, DARK_RED, 8, 4),
        ("Heading 4", 10.5, "111111", 6, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name in ("Heading 2", "Heading 4")
        style.font.italic = name == "Heading 3"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.page_break_before = False


def build_footer():
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "5")
    top.set(qn("w:color"), "D9D9D9")
    borders.append(top)
    p_pr.append(borders)
    r = p.add_run("Page | ")
    set_run_font(r, size=9)
    add_field(p, " PAGE ", "1")
    first = section.first_page_footer
    for fp in first.paragraphs:
        fp._element.getparent().remove(fp._element)
    first.add_paragraph("")


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(items, level=0):
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, BULLET_NUM_ID)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(items):
    list_num_id = create_numbering("decimal", "%1.")
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, list_num_id)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r)


def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_note(title, text, fill="FFF2CC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_fixed_cell_width(cell, 6.4)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}: ")
    set_run_font(r, bold=True, color=DARK_RED)
    r = p.add_run(text)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_cant_split(hdr)
    for idx, (cell, text) in enumerate(zip(hdr.cells, headers)):
        if widths:
            set_fixed_cell_width(cell, widths[idx])
        set_cell_shading(cell, PEACH)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=DARK_RED)
    for row_data in rows:
        row = table.add_row()
        set_cant_split(row)
        for idx, (cell, text) in enumerate(zip(row.cells, row_data)):
            if widths:
                set_fixed_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_run_font(r, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_fixed_cell_width(cell, 6.4)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 90, 130, 90, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(code)
    set_run_font(r, name="Consolas", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


figure_number = 0


def add_figure(filename, caption, width=6.35):
    global figure_number
    figure_number += 1
    path = DIAGRAMS / filename
    with Image.open(path) as im:
        w, h = im.size
    max_h = 6.8
    height = width * h / w
    if height > max_h:
        width = max_h * w / h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(f"Figure {figure_number}. {caption}")
    set_run_font(r, size=9, italic=True, color=GRAY)


def page_break():
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    r = p.add_run(" ")
    set_run_font(r, size=1)


def feature_section(number, title, overview, class_figure, class_caption, methods, seq_figure, seq_caption, query):
    page_break()
    add_heading(f"{number}. {title}", 2)
    add_body(overview)
    add_heading("a. Class/Component Diagram", 3)
    add_figure(class_figure, class_caption)
    add_heading("b. Class and Method Specifications", 3)
    add_table(
        ["Component", "Method", "Inputs / outputs and internal processing"],
        methods,
        [1.35, 1.55, 3.50],
        8.2,
    )
    add_heading("c. Sequence Diagram(s)", 3)
    add_figure(seq_figure, seq_caption)
    add_heading("d. Database Queries / Persistence Operations", 3)
    add_body("Representative operations used by the implementation are shown below. Mongoose translates these operations to MongoDB queries.")
    add_code(query)


configure_styles()
build_footer()
BULLET_NUM_ID = create_numbering("bullet", "•")

# Cover
for _ in range(3):
    doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), width=Inches(2.96))
p.paragraph_format.space_after = Pt(145)
p = doc.add_paragraph("UNI-MATE")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.runs[0]
set_run_font(r, size=30, bold=True, color=RED)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph("Software Design Specification")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.runs[0]
set_run_font(r, size=23, bold=True, color=RED)
p.paragraph_format.space_after = Pt(210)
p = doc.add_paragraph("– Ho Chi Minh City, July 2026 –")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.runs[0], size=12)

# Record of Changes
page_break()
p = doc.add_paragraph("RECORD OF CHANGES")
set_run_font(p.runs[0], size=15, bold=True, color=DARK_GREEN)
p.paragraph_format.space_after = Pt(10)
add_table(
    ["Date", "A*, M, D", "In charge", "Change Description"],
    [["14/07/2026", "A", "UNI-MATE Team", "Initial SDS baseline derived from the current React/Express/MongoDB implementation."],
     ["", "", "", ""], ["", "", "", ""]],
    [0.85, 0.70, 1.15, 3.70],
    9,
)
p = doc.add_paragraph("* A - Added; M - Modified; D - Deleted")
set_run_font(p.runs[0], size=9, italic=True)

# TOC
page_break()
p = doc.add_paragraph("Table of Contents")
set_run_font(p.runs[0], size=21, color=BLUE)
p.paragraph_format.space_after = Pt(10)
toc_p = doc.add_paragraph()
add_field(toc_p, ' TOC \\o "1-2" \\h \\z \\u ', "Right-click and update field to generate the table of contents.")

# I. Overview
page_break()
add_heading("I. Overview", 1)
add_body("UNI-MATE is a full-stack web application for students aged 18 to under 30 to discover compatible people, form mutual matches, communicate, create small groups, choose cafes, and use partner vouchers. The reviewed baseline consists of a React/Vite client, an Express/TypeScript API, MongoDB persistence through Mongoose, and Socket.IO realtime messaging.")
add_note("Implementation baseline", "This SDS documents the code present in the repository on 14 July 2026. It describes actual behavior, including the current mutual-like chat behavior and the remaining cafe proposal workflow.", LIGHT_BLUE)

add_heading("1. System Architecture and Code Packages", 2)
add_heading("a. Architectural View", 3)
add_figure("01_architecture.png", "UNI-MATE logical architecture and external integrations")
add_body("The client uses REST for transactional operations and Socket.IO for realtime chat, typing indicators, read receipts, match updates, and notifications. The backend applies JWT authentication, role checks, Zod validation, rate limiting, and centralized error handling before invoking controllers and services.")
add_heading("b. Package Dependency Diagram", 3)
add_figure("02_packages.png", "Frontend and backend package dependencies")
add_heading("c. Package Descriptions", 3)
package_rows = [
    ("01", "frontend/src/app", "Top-level React route composition and application bootstrap."),
    ("02", "frontend/src/features", "Feature screens for authentication, onboarding, discovery, match, chat, group, places, safety, partner, and admin."),
    ("03", "frontend/src/services + lib", "Axios REST client, access-token refresh interceptor, and Socket.IO client."),
    ("04", "frontend/src/stores", "Zustand authentication/session state and pending OTP context."),
    ("05", "backend/src/routes", "HTTP endpoint declarations and composition of authentication/validation middleware."),
    ("06", "backend/src/controllers", "Request orchestration, authorization-sensitive checks, response mapping, and socket emission."),
    ("07", "backend/src/services", "Authentication, compatibility scoring, distance routing, place suggestion, email, and notification logic."),
    ("08", "backend/src/models", "Mongoose schemas, constraints, references, indexes, and lifecycle fields."),
    ("09", "backend/src/sockets", "Authenticated private/group realtime messaging and read/typing events."),
    ("10", "backend/src/middlewares", "JWT authentication, admin authorization, validation, not-found, and error handling."),
    ("11", "backend/src/validations", "Zod request contracts for authentication, users, matching, messages, and safety."),
    ("12", "backend/src/utils + config", "Environment configuration, token helpers, zodiac calculation, seeding, and utility functions."),
]
add_table(["No", "Package", "Description"], package_rows, [0.45, 2.05, 3.90], 8.4)

add_heading("2. Design Conventions", 2)
add_bullets([
    "HTTP routes use resource-oriented paths under /api and JSON request/response bodies.",
    "Protected routes use Bearer access tokens; refresh tokens are also stored in an HTTP-only cookie.",
    "Mongoose ObjectId references model cross-collection relationships; selected aggregates embed arrays and value objects for read efficiency.",
    "Dates use server-side Date values and Mongoose timestamps. Geospatial coordinates use GeoJSON [longitude, latitude].",
    "Realtime rooms are named with either the ChatRoom id, group:<groupId>, or user:<userId>.",
    "Status fields implement lifecycle state rather than physical deletion for matches, rooms, groups, places, reports, and vouchers.",
])

# Database Design
page_break()
add_heading("3. Database Design", 2)
add_body("MongoDB is the system of record. Mongoose schemas enforce required fields, enumerated statuses, selected uniqueness constraints, TTL behavior for OTP records, and 2dsphere indexes for user and cafe locations.")
add_heading("a. Core Matching and Chat Schema", 3)
add_figure("03_er_core.png", "Core user discovery, match, cafe proposal, and private chat collections")
add_heading("b. Extended Social, Partner, and Moderation Schema", 3)
add_figure("04_er_extended.png", "Group, report, partner, voucher, tag, and OTP relationships")
add_heading("c. Collection Descriptions", 3)
collections = [
    ("01", "User", "Identity, role/status, profile, onboarding survey, location, preferences, block list, and token metadata."),
    ("02", "EmailOtp", "Hashed email OTP with expiry, attempt count, resend count, consumed flag, and TTL cleanup."),
    ("03", "Otp", "Legacy/simple hashed OTP record keyed by email with TTL expiry."),
    ("04", "Swipe", "Unique directed like/pass decision from one user to another."),
    ("05", "Match", "Two participants, compatibility score, lifecycle state, selected cafe, confirmations, room, and expiry."),
    ("06", "CafeProposal", "Cafe proposed for a match, proposer, cafe reference, and active/accepted/rejected history."),
    ("07", "ChatRoom", "One-to-one match room, participants, cafe, activity state, hide list, and last-message summary."),
    ("08", "Message", "Private-room text or file message, sender, type, attachment metadata, and readBy list."),
    ("09", "Notification", "Per-user notification with type, title/body, flexible data payload, and read timestamp."),
    ("10", "Group", "Owner-created group, member list, purpose, capacity, and active/dissolved state."),
    ("11", "GroupMessage", "Group text/file message with sender and readBy list."),
    ("12", "PlaceCache", "Cafe/place profile, address, geolocation, operational fields, tags, amenities, and partner ownership."),
    ("13", "Voucher", "Partner-issued cafe voucher, quota, usage count, saved users, expiry, and activation state."),
    ("14", "Report", "Reporter, reported user, optional match/room/message context, evidence, priority, and moderation result."),
    ("15", "AdminAction", "Append-only audit event with admin, action, polymorphic target type/id, and reason."),
    ("16", "Tag", "Admin-managed cafe/profile tag catalogue; places currently persist tag names as strings."),
]
add_table(["No", "Collection", "Purpose"], collections, [0.45, 1.20, 4.75], 8.1)
add_heading("d. Important Indexes and Constraints", 3)
indexes = [
    ("User", "email unique; role/status/isActive; location 2dsphere", "Authentication, admin filtering, and discovery radius queries."),
    ("Swipe", "fromUser + toUser unique", "One current directed decision per user pair."),
    ("Match", "users; status", "Participant lookup and lifecycle filtering."),
    ("ChatRoom", "match unique", "At most one room per match."),
    ("CafeProposal", "match + status", "Efficient retrieval/replacement of active proposal."),
    ("PlaceCache", "googlePlaceId unique sparse; location 2dsphere", "Cache identity and nearby cafe search."),
    ("Voucher", "code + placeId unique", "Prevents duplicate code within a cafe."),
    ("Report", "status + createdAt", "Moderation queue ordering."),
    ("EmailOtp", "email + createdAt; expiresAt TTL", "Latest challenge lookup and automatic expiry."),
]
add_table(["Collection", "Index / constraint", "Rationale"], indexes, [1.25, 2.25, 2.90], 8.2)

# II Code Designs
page_break()
add_heading("II. Code Designs", 1)
add_body("This section groups classes and runtime interactions by user-visible function. Boundary objects are React pages, control objects are controllers/services/socket handlers, and entity objects are Mongoose models.")

feature_section(
    1,
    "Authentication, Account, and Onboarding",
    "This subsystem creates accounts, validates credentials, optionally performs email OTP two-factor authentication, refreshes sessions, resets passwords, and collects the profile/preferences/location required for discovery.",
    "05_class_auth.png",
    "Authentication and profile component relationships",
    [
        ("AuthController", "registerController", "email, password -> user/tokens; delegates validation and registration, then sets refresh cookie."),
        ("AuthController", "loginController", "credentials -> tokens or requiresTwoFactor; sets refresh cookie only after full authentication."),
        ("AuthService", "registerWithPassword", "Normalizes email, rejects duplicate, hashes password, creates verified user, and issues tokens."),
        ("AuthService", "loginWithPassword", "Restores expired suspension, checks account state/password, and branches to OTP or token issue."),
        ("AuthService", "sendEmailOtp", "Generates six-digit OTP, stores bcrypt hash and expiry, enforces resend metadata, and sends email."),
        ("AuthService", "verifyEmailOtp", "Validates newest unconsumed OTP, increments failed attempts, consumes success, and issues tokens."),
        ("UserController", "completeOnboarding", "Profile/survey/location -> updated User; enforces age 18 to under 30 and computes zodiac."),
        ("UserController", "geocodeLocation", "city/district -> filtered Vietnam coordinates via Nominatim."),
        ("UserController", "updateProfile", "Whitelisted profile/preferences values -> persisted user."),
    ],
    "09_seq_login.png",
    "Password login with optional email OTP two-factor authentication",
    'const user = await User.findOne({ email });\nawait EmailOtp.create({ email, otpHash, expiresAt });\nconst otp = await EmailOtp.findOne({ email, consumed: false, expiresAt: { $gt: new Date() } })\n  .sort({ createdAt: -1 });\nawait User.findByIdAndUpdate(userId, { onboardingCompleted: true, onboarding, location });',
)

feature_section(
    2,
    "Discovery and Mutual Matching",
    "Discovery excludes prior swipes, both directions of blocking, inactive/non-user accounts, incomplete onboarding, invalid locations, incompatible age/purpose/gender filters, and users outside the configured radius. Candidates are ranked with route-time distance, cafe style, interests, goals, major preference, vibe, age, recency, and time overlap.",
    "06_class_matching_chat.png",
    "Discovery, match, cafe proposal, and chat control/entity relationships",
    [
        ("DiscoveryController", "feedController", "Authenticated user -> ranked candidate list."),
        ("DiscoveryController", "likeController", "targetUserId -> like result; emits incoming-like notification when not mutual."),
        ("MatchingService", "getDiscoveryFeed", "Loads exclusion/filter sets, geospatial candidates, OSRM route metadata, scores, sorts, and limits to 10."),
        ("MatchingService", "scoreUsers", "Two profiles and route metadata -> score, reasons, common tags/styles, distance diagnostics."),
        ("MatchingService", "swipe", "Upserts directed action; on mutual like creates/upserts Match and active ChatRoom."),
        ("MatchingService", "getIncomingLikes", "Returns users who liked the current user and have not received a response swipe."),
    ],
    "10_seq_matching.png",
    "Discovery like/pass decision and actual mutual-match behavior",
    'const swiped = await Swipe.find({ fromUser: userId }).distinct("toUser");\nconst users = await User.find({ _id: { $nin: excluded }, role: "user", status: "active", location: geoFilter });\nawait Swipe.findOneAndUpdate({ fromUser, toUser }, { action }, { upsert: true, new: true });\nconst reciprocal = await Swipe.findOne({ fromUser: target, toUser: actor, action: "like" });\nawait Match.create({ users: [actor, target], status: "chat_opened", expiresAt });',
)
add_note("Actual behavior", "Mutual like currently creates an active ChatRoom immediately. This conflicts with the README rule that chat should remain locked until both users confirm the same cafe.", "F4CCCC")

feature_section(
    3,
    "Cafe Proposal and Match Lifecycle",
    "A participant can retrieve suggested cafes, replace the active proposal, and notify the other participant. The proposal creator is automatically recorded as the first confirmer. Only the other participant can confirm or reject. Confirmation upserts an active room; rejection clears the selection; cancellation archives an existing room.",
    "06_class_matching_chat.png",
    "MatchController relationships to Match, CafeProposal, PlaceCache, and ChatRoom",
    [
        ("MatchController", "placeSuggestions", "matchId -> suggested active cafes; rejects missing/expired match."),
        ("MatchController", "selectPlace", "matchId/placeId -> proposal; replaces active proposal and sets cafe_proposed."),
        ("MatchController", "confirmPlace", "Adds non-proposer to confirmedBy, accepts proposal, upserts room, and sets chat_opened."),
        ("MatchController", "rejectPlace", "Non-proposer rejects; clears selection/confirmations and returns match to matched."),
        ("MatchController", "cancelMatch", "Sets cancelled, rejects active proposals, archives room, and notifies peer."),
        ("PlacesService", "suggestCafePlaces", "Uses participant locations and cached/seeded places to return cafe candidates."),
    ],
    "11_seq_cafe.png",
    "Cafe selection, confirmation, rejection, and realtime match update",
    'await CafeProposal.updateMany({ match: match._id, status: "active" }, { status: "replaced" });\nawait CafeProposal.create({ match: match._id, proposedBy: userId, cafe: placeId, status: "active" });\nawait Match.findOneAndUpdate({ _id: matchId }, { $addToSet: { confirmedBy: userId } }, { new: true });\nawait ChatRoom.findOneAndUpdate({ match: matchId }, roomData, { upsert: true, new: true });',
)

feature_section(
    4,
    "Private Chat and Notifications",
    "Private chat supports room listing/hiding, chat history, text and file messages, typing indicators, delivery broadcasts, read receipts, and per-recipient notifications. Both HTTP and Socket.IO message submission paths enforce active room and chat_opened match state.",
    "06_class_matching_chat.png",
    "ChatController and SocketHandler relationships to room and message entities",
    [
        ("ChatController", "listRooms", "Returns rooms containing user and not hidden by that user, ordered by lastMessageAt."),
        ("ChatController", "getRoom", "Verifies membership and states, calculates bilateral block state, and returns up to 100 messages."),
        ("ChatController", "sendMessage", "HTTP fallback: validates message, persists it, updates room, broadcasts, and notifies peer."),
        ("SocketHandler", "join_room", "Joins room only when the socket user is a participant and room is active."),
        ("SocketHandler", "send_message", "Realtime validation/persistence/broadcast plus notification creation."),
        ("SocketHandler", "mark_read", "Adds reader id to readBy for all room messages and emits message_read."),
        ("NotificationService", "createAndEmitNotification", "Persists Notification then emits notification:new to user:<id>."),
    ],
    "12_seq_chat.png",
    "Realtime private message, notification, and read-receipt flow",
    'const room = await ChatRoom.findOne({ _id: roomId, users: userId });\nconst match = await Match.findById(room.match);\nconst message = await Message.create({ room: room._id, sender: userId, text, type, fileUrl, readBy: [userId] });\nawait Message.updateMany({ room: roomId }, { $addToSet: { readBy: userId } });',
)

feature_section(
    5,
    "Groups and Group Chat",
    "Users can create purpose-oriented groups, add members by email, remove members or leave, dissolve a group, and exchange realtime or HTTP-fallback group messages. Owner, capacity, membership, and active-state rules are enforced in controllers/socket handlers.",
    "07_class_group_partner.png",
    "Group management and group-chat components",
    [
        ("GroupController", "createGroup", "Creates group with requester as creator and first member; default capacity is six."),
        ("GroupController", "getMyGroups", "Lists active groups containing the requester and populates creator/members."),
        ("GroupController", "addMember", "Owner-only; validates capacity and active target user found by email."),
        ("GroupController", "removeMember", "Owner may remove others; member may leave; owner cannot leave without dissolving."),
        ("GroupController", "dissolveGroup", "Owner-only transition to dissolved."),
        ("GroupChatController", "getGroupMessages", "Verifies member and returns group history."),
        ("SocketHandler", "send_group_message", "Validates active membership, persists GroupMessage, broadcasts, and notifies peers."),
    ],
    "13_seq_group.png",
    "Group creation, member addition, and realtime group messaging",
    'const group = await Group.create({ name, creator: userId, members: [userId], purpose, maxMembers });\nconst target = await User.findOne({ email: normalizedEmail, status: "active" });\ngroup.members.push(target._id); await group.save();\nconst messages = await GroupMessage.find({ group: groupId }).sort({ createdAt: 1 });',
)

feature_section(
    6,
    "Places, Partner Registration, and Vouchers",
    "Authenticated users browse active cafes and available vouchers. A prospective owner submits one pending partner-place application; admin approval activates the place and upgrades the owner role to partner. Partners manage owned place data and voucher lifecycle. Users save valid vouchers atomically against quota.",
    "07_class_group_partner.png",
    "Places, partner ownership, and voucher component relationships",
    [
        ("PlacesController", "listPlaces/getPlace", "Returns only active place records."),
        ("PlacesController", "registerPartnerPlace", "Validates cafe form/hours and creates one pending partner place per applicant."),
        ("PlacesController", "getPlaceVouchers", "Returns active, unexpired, non-exhausted vouchers and savedByMe flag."),
        ("PlacesController", "saveVoucher", "Idempotent save; atomically adds user and increments usage only while valid."),
        ("PartnerController", "getMyPlaces/updateMyPlace", "Lists owned places and updates whitelisted fields after ownership check."),
        ("PartnerController", "createVoucher", "Requires partner ownership, active cafe, required values, and unique code within cafe."),
        ("PartnerController", "toggleVoucherStatus/deleteVoucher", "Owner-scoped voucher activation toggle or deletion."),
        ("AdminController", "hidePlace", "Transitions place status, upgrades partner role on approval, notifies owner, and audits."),
    ],
    "15_seq_partner_voucher.png",
    "Partner place approval, voucher creation, and user save flow",
    'const place = await PlaceCache.create({ ...data, partnerId: userId, isPartnerPlace: true, status: "pending" });\nconst existing = await Voucher.findOne({ placeId, code: code.toUpperCase() });\nconst voucher = await Voucher.findOneAndUpdate(validVoucherFilter, { $addToSet: { savedBy: userId }, $inc: { currentUsageCount: 1 } }, { new: true });',
)

feature_section(
    7,
    "Safety, Blocking, Reporting, and Moderation",
    "Safety operations validate report context, prioritize repeated independent reports, block communication across related matches/rooms, and restore communication only when neither party continues blocking. Admin moderation supports dismiss, warn, suspend, and ban with warning-count escalation and audit logging.",
    "08_class_safety_admin.png",
    "Safety and moderation control/entity relationships",
    [
        ("SafetyController", "reportUser", "Creates report with optional match/room/message/evidence and marks priority after three distinct reporters."),
        ("SafetyController", "blockUser", "Adds target to blockedUsers and sets all pair Matches/ChatRooms to blocked."),
        ("SafetyController", "unblockUser", "Removes block; restores chat_opened/active only if target does not still block requester."),
        ("AdminController", "adminReportDetail", "Loads report context and up to 200 room messages for investigation."),
        ("AdminController", "updateReport", "Applies one terminal action, updates account/report, emits notice, disconnects locked user."),
        ("AdminController", "audit", "Creates AdminAction with admin/action/targetType/targetId/reason."),
    ],
    "14_seq_safety_admin.png",
    "Report submission and admin moderation decision",
    'const report = await Report.create({ reporter, reportedUser, match, room, message, reason, evidenceUrls });\nconst reporters = await Report.distinct("reporter", { reportedUser });\nawait User.findByIdAndUpdate(actor, { $addToSet: { blockedUsers: target } });\nawait Match.updateMany({ users: { $all: [actor, target] } }, { status: "blocked" });\nawait AdminAction.create({ admin, action, targetType, targetId, reason });',
)

feature_section(
    8,
    "Administration and Operational Oversight",
    "Admin pages provide dashboard counts, user search/detail/update/status control, report queue/detail/resolution, match search, place review/update/delete/status control, and audit log listing. All admin routes require authentication and the admin role.",
    "08_class_safety_admin.png",
    "Administration controller dependencies and audited targets",
    [
        ("AdminController", "dashboard", "Parallel counts for users, new users, matches, confirmed/opened matches, new reports, and active places."),
        ("AdminController", "adminUsers/adminUserDetail", "Filters users and loads selected user's match/report history."),
        ("AdminController", "createAdminUser/updateAdminUser", "Admin-managed account creation and full editable profile/role/status update."),
        ("AdminController", "updateUserStatus", "Transitions active/suspended/banned and records audit reason."),
        ("AdminController", "adminMatches", "Filters by status or searches participant, cafe, object id, and status text."),
        ("AdminController", "adminPlaces/upsertPlace/hidePlace/deletePlace", "Review and lifecycle management of cafe records."),
        ("AdminController", "adminActions", "Returns latest 100 populated audit events."),
    ],
    "14_seq_safety_admin.png",
    "Representative admin review and audited moderation flow",
    'const [users, matches, reports, places] = await Promise.all([User.countDocuments(), Match.countDocuments(), Report.countDocuments({ status: "new" }), PlaceCache.countDocuments({ status: "active" })]);\nconst result = await User.findByIdAndUpdate(userId, statusUpdate, { new: true });\nconst actions = await AdminAction.find().populate("admin", "email displayName").sort({ createdAt: -1 }).limit(100);',
)

# III Cross-cutting design
page_break()
add_heading("III. Cross-Cutting Design", 1)
add_heading("1. Security and Authorization", 2)
security = [
    ("Authentication", "Short-lived JWT access token; refresh token hash persisted and HTTP-only refresh cookie supported."),
    ("Authorization", "requireAuth on protected routers; requireAdmin and requirePartner on privileged areas; ownership/member checks inside controllers."),
    ("Password/OTP", "bcrypt password and OTP hashes; expiring/consumable email OTP; attempt limit; admin always requires 2FA."),
    ("Input validation", "Zod request validation; Mongoose enum/range constraints; upload type/size checks."),
    ("HTTP hardening", "Helmet, CORS with configured client origin, 2 MB JSON limit, cookie parser, and 500 requests per 15-minute rate limit."),
    ("Account enforcement", "Suspended/banned/inactive users cannot log in; moderated sockets are disconnected."),
    ("Privacy", "Password and refresh hashes excluded from normal/admin projections; report context restricted to authorized admin routes."),
]
add_table(["Concern", "Current design"], security, [1.45, 4.95], 8.5)

add_heading("2. Realtime Design", 2)
events = [
    ("join_room", "Client -> Server", "Verify private-room membership and active state, then join room id."),
    ("send_message", "Client -> Server", "Persist and broadcast private message; notify other participant."),
    ("new_message", "Server -> Room", "Deliver populated private message."),
    ("typing_start/stop", "Bidirectional", "Broadcast transient private typing state."),
    ("mark_read/message_read", "Bidirectional", "Persist readBy and inform peer."),
    ("join_group", "Client -> Server", "Verify active group membership and join group:<id>."),
    ("send_group_message", "Client -> Server", "Persist/broadcast group message and notify other members."),
    ("new_group_message", "Server -> Group", "Deliver populated group message."),
    ("group_typing_start/stop", "Bidirectional", "Broadcast group typing state."),
    ("mark_group_read", "Client -> Server", "Persist group readBy and emit receipt."),
    ("notification:new", "Server -> User", "Push persisted notification to user:<id>."),
    ("match:updated", "Server -> Users", "Push cafe/match/chat lifecycle change."),
]
add_table(["Event", "Direction", "Responsibility"], events, [1.75, 1.15, 3.50], 8.2)

add_heading("3. State Models", 2)
states = [
    ("Match", "matched -> cafe_proposed -> chat_opened; any -> expired/blocked/cancelled", "Mutual like currently creates chat_opened directly."),
    ("CafeProposal", "active -> accepted/replaced/expired/rejected", "Only one active proposal is intended per match."),
    ("ChatRoom", "active <-> blocked; active -> archived", "Cancel archives; bilateral unblock may reactivate."),
    ("Group", "active -> dissolved", "Dissolved groups reject joins/messages."),
    ("PlaceCache", "pending -> active/hidden/rejected", "Active partner place upgrades owner to partner."),
    ("Report", "new -> reviewing -> resolved_valid/resolved_invalid/dismissed", "Controller currently directly uses dismissed/resolved_valid for final actions."),
    ("User", "active -> suspended/banned; suspended -> active after expiry", "Warning policy can escalate automatically."),
]
add_table(["Aggregate", "Transitions", "Notes"], states, [1.15, 2.85, 2.40], 8.2)

add_heading("4. External Services and Failure Handling", 2)
external = [
    ("OSRM", "Travel-time matrix for compatibility scoring.", "Failure is logged; distance score falls back to zero."),
    ("Google Places / cache", "Cafe suggestions and metadata.", "Cached or seeded fallback places keep the flow testable."),
    ("Nominatim", "Manual district/city geocoding.", "Non-2xx returns 502; results are filtered to Vietnam/city."),
    ("SMTP/email", "OTP delivery.", "Development may print OTP when configured; service errors propagate through centralized handler."),
    ("Local uploads", "Chat files and report evidence.", "20 MB chat limit; 5 MB image-only report evidence."),
]
add_table(["Dependency", "Use", "Fallback / control"], external, [1.45, 2.35, 2.60], 8.3)

# Appendices
page_break()
add_heading("Appendix A. REST API Inventory", 1)
add_body("All routes below are prefixed with /api. Except for register/login/OTP/refresh, routes require an authenticated access token. Admin and partner groups apply additional role checks.")
api_rows = [
    ("Auth", "POST", "/auth/register", "Create account and issue tokens."),
    ("Auth", "POST", "/auth/login", "Password login; may require 2FA."),
    ("Auth", "POST", "/auth/send-otp", "Send email OTP."),
    ("Auth", "POST", "/auth/verify-otp", "Verify OTP and issue tokens."),
    ("Auth", "POST", "/auth/forgot-password/send-otp", "Start reset."),
    ("Auth", "POST", "/auth/forgot-password/verify-otp", "Return reset token."),
    ("Auth", "POST", "/auth/forgot-password/reset", "Set new password."),
    ("Auth", "POST", "/auth/refresh", "Rotate/refresh session."),
    ("Auth", "GET/POST", "/auth/me; /auth/logout", "Current user; terminate session."),
    ("Users", "GET/PATCH", "/users/profile", "Read/update profile."),
    ("Users", "PUT/PATCH", "/users/photos; /users/password", "Photos and password."),
    ("Users", "DELETE", "/users/profile", "Delete/deactivate account."),
    ("Users", "POST", "/users/onboarding", "Complete profile/survey/location."),
    ("Users", "POST/PATCH", "/users/location/geocode; /users/location", "Geocode/update location."),
    ("Users", "POST", "/users/avatar; /users/photo", "Upload images."),
    ("Discovery", "GET", "/discovery; /discovery/incoming-likes", "Feed and pending incoming likes."),
    ("Discovery", "POST", "/discovery/like; /discovery/pass", "Persist swipe and possibly match."),
    ("Matches", "GET", "/matches; /matches/:id", "List/detail."),
    ("Matches", "GET", "/matches/:id/place-suggestions", "Suggested cafes."),
    ("Matches", "POST", "/matches/:id/select-place", "Create/replace proposal."),
    ("Matches", "POST", "/matches/:id/confirm-place", "Accept proposal."),
    ("Matches", "POST", "/matches/:id/reject-place; /cancel", "Reject cafe or cancel match."),
    ("Chat", "GET", "/chat; /chat/:roomId", "Room list/detail/history."),
    ("Chat", "POST/DELETE", "/chat/:roomId/messages; /chat/:roomId", "Send via HTTP; hide conversation."),
    ("Notifications", "GET/PATCH", "/notifications; /read-all; /:id/read", "List and mark read."),
    ("Groups", "GET/POST", "/groups", "List/create groups."),
    ("Groups", "GET", "/groups/:id; /groups/:id/messages", "Group detail/history."),
    ("Groups", "POST/DELETE", "/groups/:id/members; /members/:userId", "Add/remove/leave."),
    ("Groups", "POST", "/groups/:id/dissolve; /groups/:id/messages", "Dissolve/send fallback."),
    ("Places", "GET", "/places; /places/:id; /places/:id/vouchers", "Browse active places/vouchers."),
    ("Places", "POST", "/places/partner-register", "Submit partner cafe."),
    ("Places", "GET", "/places/my-partner-registration; /saved-vouchers", "Application/saved vouchers."),
    ("Places", "POST/PATCH", "/places/:placeId/vouchers/:voucherId/save", "Save voucher."),
    ("Partner", "GET/PATCH", "/partner/places; /places/:placeId", "Owned cafes."),
    ("Partner", "GET/POST", "/partner/places/:placeId/vouchers", "List/create voucher."),
    ("Partner", "PATCH/DELETE", "/partner/vouchers/:id/toggle; /:id", "Toggle/delete voucher."),
    ("Safety", "POST", "/safety/report; /block; /unblock", "Report and bilateral safety controls."),
    ("Upload", "POST", "/upload/chat-file; /report-evidence", "File/evidence uploads."),
    ("Admin", "GET", "/admin/dashboard; /analytics; /settings", "Metrics and settings view."),
    ("Admin", "GET/POST/PUT/PATCH", "/admin/users[/:id][/status]", "User administration."),
    ("Admin", "GET/PATCH", "/admin/reports[/:id]", "Report review/resolution."),
    ("Admin", "GET", "/admin/matches", "Match oversight."),
    ("Admin", "GET/PUT/PATCH/DELETE", "/admin/places[/:id][/status]", "Place review/maintenance."),
    ("Admin", "GET", "/admin/actions", "Audit log."),
]
add_table(["Area", "Method", "Path", "Purpose"], api_rows, [0.85, 0.75, 2.60, 2.20], 7.6)

page_break()
add_heading("Appendix B. Configuration and Deployment", 1)
add_heading("1. Runtime Stack", 2)
add_numbered([
    "Frontend: React, TypeScript, Vite, React Router, Zustand, Axios, Tailwind CSS, Socket.IO Client.",
    "Backend: Node.js, Express, TypeScript, Mongoose, JWT, bcryptjs, Zod, Socket.IO, Multer, Nodemailer.",
    "Database: MongoDB with geospatial and TTL indexes.",
    "Development entry points: frontend http://localhost:5173 and backend http://localhost:5000.",
])
add_heading("2. Required Configuration Categories", 2)
config_rows = [
    ("Database", "MongoDB connection URI."),
    ("Authentication", "JWT access/refresh secrets and token lifetimes; cookie secure flag."),
    ("Client", "Allowed client origin / frontend URL."),
    ("Email", "SMTP host, port, credentials, sender and development OTP logging option."),
    ("Places", "Google Maps/Places API key or cached/seeded fallback."),
    ("Distance", "OSRM endpoint and scoring thresholds/weights/candidate limit."),
    ("Uploads", "Writable uploads directory and public /uploads route."),
]
add_table(["Category", "Configuration responsibility"], config_rows, [1.55, 4.85], 8.6)

add_heading("3. Build and Run", 2)
add_numbered([
    "Provide backend and frontend environment files.",
    "Install backend dependencies and start MongoDB or configure MongoDB Atlas.",
    "Run the backend development process.",
    "Install frontend dependencies and run the Vite client.",
    "Verify /health, authentication, protected REST routes, and Socket.IO connectivity.",
])

page_break()
add_heading("Appendix C. Known Implementation Discrepancies and Risks", 1)
risks = [
    ("High", "Cafe gate bypass", "Mutual like creates Match(status=chat_opened) and active ChatRoom immediately, contrary to the README and admin setting cafeGateRequired=true."),
    ("High", "State regression on cafe selection", "Selecting a cafe after mutual like changes chat_opened back to cafe_proposed, temporarily locking an existing chat."),
    ("Medium", "Registration verification", "Registration sets emailVerified=true without OTP; OTP is only used for 2FA/login and password reset."),
    ("Medium", "Match participant cardinality", "Match.users is an unrestricted array; schema does not enforce exactly two users."),
    ("Medium", "Tag consistency", "Tag has its own catalogue collection, but PlaceCache stores tag names as strings and does not reference Tag ids."),
    ("Medium", "Unblock restoration", "Unblock restores all blocked pair matches to chat_opened, even if their prior state was not an open chat."),
    ("Low", "Dual OTP models", "Both Otp and EmailOtp exist; active auth service primarily uses EmailOtp."),
    ("Low", "Local file storage", "Uploads are stored on the application host and require persistent/shared storage for scaled deployment."),
]
add_table(["Severity", "Issue", "Design impact"], risks, [0.75, 1.55, 4.10], 8.2)
add_note("Recommended decision", "Either enforce cafe confirmation before creating ChatRoom, or formally remove the cafe-gated requirement and stop moving chat_opened matches back to cafe_proposed. The current hybrid state should not be treated as a stable business rule.", "F4CCCC")

# Ensure Word refreshes fields on open even before the COM update step.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "UNI-MATE Software Design Specification"
doc.core_properties.subject = "Implementation-derived software design specification"
doc.core_properties.author = "UNI-MATE Team"
doc.core_properties.keywords = "UNI-MATE, SDS, UML, React, Express, MongoDB"
doc.save(OUTPUT)
print(OUTPUT)
